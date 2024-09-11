import os
import re
from docx import Document
import fitz  # PyMuPDF for PDF files

# Caches for file content and normalized text
file_cache = {}
text_cache = {}

def is_machine_readable(file_path):
    """Determine if the file is machine-readable by attempting to open and read it."""
    if file_path.lower().endswith('.pdf'):
        try:
            doc = fitz.open(file_path)
            doc.close()
            return True
        except:
            return False
    elif file_path.lower().endswith('.docx'):
        try:
            doc = Document(file_path)
            return True
        except:
            return False
    return False

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file using PyMuPDF."""
    if file_path in file_cache:
        return file_cache[file_path]
    
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        file_cache[file_path] = text
        return text
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return ""

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file using python-docx."""
    if file_path in file_cache:
        return file_cache[file_path]
    
    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        file_cache[file_path] = text
        return text
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return ""

def normalize_text(text):
    """Remove extra spaces, newlines, and normalize the text for matching."""
    return ' '.join(text.split())

def search_files(search_term, directory):
    found_files = set()  # Use a set to avoid duplicates
    search_pattern = re.compile(re.escape(search_term), re.IGNORECASE)  # Case-insensitive search

    for root, dirs, files in os.walk(directory):
        for file in files:
            file_path = os.path.join(root, file)
            if not is_machine_readable(file_path):
                continue

            # Extract text based on file type
            if file.lower().endswith('.pdf'):
                text = extract_text_from_pdf(file_path)
            elif file.lower().endswith('.docx'):
                text = extract_text_from_docx(file_path)
            else:
                continue

            # Use text cache to avoid redundant normalization
            if file_path not in text_cache:
                normalized_text = normalize_text(text)
                text_cache[file_path] = normalized_text
            else:
                normalized_text = text_cache[file_path]

            # Check if the normalized search term is in the normalized text
            if search_pattern.search(normalized_text):
                found_files.add(file_path)  # Add file path to set to avoid duplicates

    # Output results
    return [os.path.basename(file_path) for file_path in found_files]

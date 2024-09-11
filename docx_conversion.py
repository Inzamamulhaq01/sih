import docx2txt
from PIL import Image
import pytesseract
from docx import Document
import os

# Path to Tesseract-OCR executable (Update if necessary)
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'


def is_scanned_word(file_path):
    
    try:
        doc = Document(file_path)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():  # Check if paragraph has any text
                text.append(para.text.strip())
        
        return not bool(text)
    
    except Exception as e:
        # Log or handle errors as needed
        print(f"Error processing document: {e}")
        return True  # Assuming it's scanned if we can't process it


def convert_docx_to_searchable(docx_path, output_path):
    # Extract images from the DOCX file
    temp_dir = "temp_images"
    os.makedirs(temp_dir, exist_ok=True)
    docx2txt.process(docx_path, temp_dir)

    # Create a new Word document for the searchable text
    doc = Document()

    # Iterate through extracted images and perform OCR
    for image_file in os.listdir(temp_dir):
        image_path = os.path.join(temp_dir, image_file)
        if image_file.endswith(('png', 'jpg', 'jpeg')):
            img = Image.open(image_path)
            text = pytesseract.image_to_string(img)
            doc.add_paragraph(text)

    # Save the searchable Word document
    doc.save(output_path)
    print(f"Searchable Word document saved to {output_path}")

    # Clean up temporary image files
    for image_file in os.listdir(temp_dir):
        os.remove(os.path.join(temp_dir, image_file))
    os.rmdir(temp_dir)
